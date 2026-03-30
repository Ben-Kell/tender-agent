from pathlib import Path
from typing import Dict, Any, List

from docx import Document
from pypdf import PdfReader


SUPPORTED_EXTENSIONS = {
    ".md",
    ".txt",
    ".pdf",
    ".docx",
}


def normalise_filename(source_name: str) -> str:
    source_path = Path(source_name)
    return f"{source_path.stem}.md"


def clean_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    lines = [line.rstrip() for line in text.splitlines()]

    cleaned_lines: List[str] = []
    previous_blank = False

    for line in lines:
        is_blank = not line.strip()
        if is_blank and previous_blank:
            continue
        cleaned_lines.append(line)
        previous_blank = is_blank

    return "\n".join(cleaned_lines).strip()


def wrap_with_header(file_path: Path, document_type: str, body: str) -> str:
    body = clean_text(body)

    return f"""# Source: {file_path.name}
# Document Type: {document_type}
# Normalised For: Tender agent ingestion

{body}
"""


def normalise_text_file(file_path: Path) -> str:
    body = file_path.read_text(encoding="utf-8", errors="ignore")
    return wrap_with_header(file_path, "txt", body)


def normalise_markdown_file(file_path: Path) -> str:
    body = file_path.read_text(encoding="utf-8", errors="ignore")
    return wrap_with_header(file_path, "md", body)


def table_to_markdown(table) -> str:
    rows = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            text = clean_text(cell.text).replace("\n", " ").strip()
            cells.append(text if text else " ")
        rows.append(cells)

    if not rows:
        return ""

    max_cols = max(len(r) for r in rows)
    padded_rows = [r + [" "] * (max_cols - len(r)) for r in rows]

    header = padded_rows[0]
    separator = ["---"] * max_cols
    body_rows = padded_rows[1:]

    markdown_lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(separator) + " |",
    ]

    for row in body_rows:
        markdown_lines.append("| " + " | ".join(row) + " |")

    return "\n".join(markdown_lines)


def normalise_docx(file_path: Path) -> str:
    doc = Document(str(file_path))

    parts: List[str] = []

    # Paragraphs
    for paragraph in doc.paragraphs:
        text = clean_text(paragraph.text)
        if text:
            parts.append(text)

    # Tables
    for table in doc.tables:
        md_table = table_to_markdown(table)
        if md_table:
            parts.append(md_table)

    body = "\n\n".join(parts).strip()

    if not body:
        body = "[No extractable text found in DOCX.]"

    return wrap_with_header(file_path, "docx", body)


def normalise_pdf(file_path: Path) -> str:
    reader = PdfReader(str(file_path))
    page_chunks: List[str] = []

    for page_number, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""

        text = clean_text(text)

        if text:
            page_chunks.append(f"## Page {page_number}\n\n{text}")
        else:
            page_chunks.append(f"## Page {page_number}\n\n[No extractable text found on this page.]")

    body = "\n\n".join(page_chunks).strip()

    if not body:
        body = "[No extractable text found in PDF.]"

    return wrap_with_header(file_path, "pdf", body)


def normalise_one_document(file_path: Path) -> str:
    suffix = file_path.suffix.lower()

    if suffix == ".txt":
        return normalise_text_file(file_path)

    if suffix == ".md":
        return normalise_markdown_file(file_path)

    if suffix == ".pdf":
        return normalise_pdf(file_path)

    if suffix == ".docx":
        return normalise_docx(file_path)

    raise ValueError(f"Unsupported file type: {file_path.suffix}")


def normalise_all_documents(tender_id: str) -> Dict[str, Any]:
    tender_id = tender_id.strip()
    if not tender_id:
        raise ValueError("tender_id is required")

    source_dir = Path("tenders") / tender_id / "input" / "01_customer_issued"
    output_dir = Path("tenders") / tender_id / "input" / "02_normalised"
    output_dir.mkdir(parents=True, exist_ok=True)

    processed: List[str] = []
    skipped: List[str] = []
    failed: List[Dict[str, str]] = []

    for file_path in source_dir.iterdir():
        if not file_path.is_file():
            continue

        if file_path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            skipped.append(str(file_path))
            continue

        output_file = output_dir / normalise_filename(file_path.name)

        try:
            normalised_content = normalise_one_document(file_path)
            output_file.write_text(normalised_content, encoding="utf-8")
            processed.append(str(output_file))
        except Exception as e:
            failed.append({
                "source_file": str(file_path),
                "error": str(e),
            })

    return {
        "source_dir": str(source_dir),
        "output_dir": str(output_dir),
        "processed": processed,
        "skipped": skipped,
        "failed": failed,
    }